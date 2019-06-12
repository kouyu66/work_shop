from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
def word_setting(document):
    # title confige
    document.styles['Title'].font.name = 'Cambria'
    document.styles['Title'].font.size = Pt(28)
    document.styles['Title'].font.bold = True
    document.styles['Title'].font.color.rgb = RGBColor(0,0,0)
    document.styles['Title'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Heading 1 config
    document.styles['Heading 1'].font.name = 'Calibri'
    document.styles['Heading 1'].font.size = Pt(18)
    document.styles['Heading 1'].font.bold = True
    document.styles['Heading 1'].font.color.rgb = RGBColor(0,0,0)

    return document
# 加入大标题
def add_title(product_name_type, document):
    title = 'SSD Test Report for {0}'.format(product_name_type)
    document.add_heading(title,0)
    return
# 加入一级标题
def add_sub_title(txt, document):
    document.add_heading(txt, level=1)
    return

# ------ Test Zone ------ #

product_name_type = 'PM983 U.2'
environment = 'Test Environment'

# 创建文档
document = Document()
# 初始化
document = word_setting(document)
add_title(product_name_type, document)
add_sub_title(environment, document)
document.save('test.docx')